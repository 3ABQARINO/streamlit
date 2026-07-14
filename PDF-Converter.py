import io
import os
import re
import zipfile
from pathlib import Path

import pandas as pd
import pdfplumber
import streamlit as st
from docx import Document


st.set_page_config(
    page_title="PDF Converter Pro",
    page_icon="📄",
    layout="wide"
)

st.title("📄 PDF Converter Pro")
st.caption("Upload one or more PDF files, extract text and tables, preview the results, and export to Word, Excel, CSV, or a ZIP package.")


# -----------------------------
# Helpers
# -----------------------------
def clean_sheet_name(name: str, used_names: set) -> str:
    cleaned = re.sub(r"[\\/*?:\[\]]", "_", name)[:31] or "Sheet"
    candidate = cleaned
    i = 1
    while candidate in used_names:
        suffix = f"_{i}"
        candidate = (cleaned[: 31 - len(suffix)] + suffix)[:31]
        i += 1
    used_names.add(candidate)
    return candidate


@st.cache_data(show_spinner=False)
def analyze_pdf(pdf_bytes: bytes, file_name: str, table_mode: str):
    text_parts = []
    table_items = []
    page_summaries = []

    table_settings = None
    if table_mode == "Borderless/text-aligned tables":
        table_settings = {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
            "intersection_tolerance": 5,
        }
    elif table_mode == "Mixed/try harder":
        table_settings = {
            "vertical_strategy": "lines",
            "horizontal_strategy": "text",
            "snap_tolerance": 4,
            "join_tolerance": 4,
            "intersection_tolerance": 4,
        }

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total_pages = len(pdf.pages)

        for page_no, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            text_parts.append(f"\n\n--- Page {page_no} ---\n\n{page_text}")

            try:
                page_tables = page.extract_tables(table_settings=table_settings)
            except TypeError:
                page_tables = page.extract_tables()

            valid_tables = []
            for table_index, table in enumerate(page_tables or [], start=1):
                if not table:
                    continue
                max_len = max(len(row) if row else 0 for row in table)
                normalized = []
                for row in table:
                    row = row or []
                    normalized.append(list(row) + [None] * (max_len - len(row)))
                if len(normalized) == 1:
                    df = pd.DataFrame(normalized)
                else:
                    header = [str(x).strip() if x is not None and str(x).strip() else f"column_{i+1}" for i, x in enumerate(normalized[0])]
                    df = pd.DataFrame(normalized[1:], columns=header)
                df.insert(0, "table_number", table_index)
                df.insert(0, "page_number", page_no)
                valid_tables.append(df)
                table_items.append({
                    "page_number": page_no,
                    "table_number": table_index,
                    "dataframe": df,
                })

            page_summaries.append({
                "page_number": page_no,
                "text_characters": len(page_text),
                "tables_found": len(valid_tables),
                "has_text": bool(page_text.strip()),
            })

    full_text = "".join(text_parts).strip()
    page_summary_df = pd.DataFrame(page_summaries)

    return {
        "file_name": file_name,
        "total_pages": total_pages,
        "full_text": full_text,
        "tables": table_items,
        "page_summary_df": page_summary_df,
        "total_tables": len(table_items),
        "has_text": bool(full_text.strip()),
    }


def make_docx_bytes(title: str, text: str, page_summary_df: pd.DataFrame | None = None) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    if page_summary_df is not None and not page_summary_df.empty:
        doc.add_paragraph(
            f"Pages: {len(page_summary_df)} | Pages with detected text: {int(page_summary_df['has_text'].sum())} | Tables detected: {int(page_summary_df['tables_found'].sum())}"
        )
    if text.strip():
        for chunk in text.split("\n\n"):
            if chunk.strip():
                doc.add_paragraph(chunk.strip())
    else:
        doc.add_paragraph("No text could be extracted. This PDF may be scanned or image-based.")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()



def make_excel_bytes(file_name: str, tables: list[pd.DataFrame], page_summary_df: pd.DataFrame | None = None) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        used = set()
        if page_summary_df is not None and not page_summary_df.empty:
            page_summary_df.to_excel(writer, index=False, sheet_name=clean_sheet_name("summary", used))

        if tables:
            all_tables = pd.concat(tables, ignore_index=True, sort=False)
            all_tables.to_excel(writer, index=False, sheet_name=clean_sheet_name("all_tables", used))

            for idx, df in enumerate(tables, start=1):
                page_no = df["page_number"].iloc[0] if "page_number" in df.columns and not df.empty else idx
                table_no = df["table_number"].iloc[0] if "table_number" in df.columns and not df.empty else idx
                sheet_name = clean_sheet_name(f"p{page_no}_t{table_no}", used)
                df.to_excel(writer, index=False, sheet_name=sheet_name)
        else:
            pd.DataFrame({"message": [f"No tables detected in {file_name}"]}).to_excel(
                writer, index=False, sheet_name=clean_sheet_name("no_tables", used)
            )
    return bio.getvalue()



def make_zip_bytes(results: list[dict], include_docx: bool, include_excel: bool, include_csv: bool, include_txt: bool) -> bytes:
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for result in results:
            stem = Path(result["file_name"]).stem

            if include_txt:
                txt_data = result["full_text"] if result["full_text"].strip() else "No text could be extracted."
                zf.writestr(f"{stem}/{stem}_extracted_text.txt", txt_data)

            if include_docx:
                docx_bytes = make_docx_bytes(result["file_name"], result["full_text"], result["page_summary_df"])
                zf.writestr(f"{stem}/{stem}_text.docx", docx_bytes)

            if include_excel:
                excel_bytes = make_excel_bytes(
                    result["file_name"],
                    [t["dataframe"] for t in result["tables"]],
                    result["page_summary_df"],
                )
                zf.writestr(f"{stem}/{stem}_tables.xlsx", excel_bytes)

            if include_csv:
                if result["tables"]:
                    for t in result["tables"]:
                        df = t["dataframe"]
                        page_no = t["page_number"]
                        table_no = t["table_number"]
                        zf.writestr(
                            f"{stem}/csv/{stem}_page_{page_no}_table_{table_no}.csv",
                            df.to_csv(index=False),
                        )
                else:
                    zf.writestr(f"{stem}/csv/{stem}_no_tables.csv", "message\nNo tables detected\n")

            zf.writestr(f"{stem}/{stem}_page_summary.csv", result["page_summary_df"].to_csv(index=False))

    return bio.getvalue()


# -----------------------------
# Sidebar
# -----------------------------
st.sidebar.header("Upload")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDF file(s)",
    type=["pdf"],
    accept_multiple_files=True,
)

table_mode = st.sidebar.selectbox(
    "Table extraction mode",
    [
        "Standard lined tables",
        "Borderless/text-aligned tables",
        "Mixed/try harder",
    ],
    index=0,
)

export_docx = st.sidebar.checkbox("Enable Word (.docx) export", value=True)
export_excel = st.sidebar.checkbox("Enable Excel (.xlsx) export", value=True)
export_csv = st.sidebar.checkbox("Enable CSV export", value=True)
export_txt = st.sidebar.checkbox("Enable TXT export", value=True)

if not uploaded_files:
    st.info("Upload one or more PDF files from the sidebar to begin.")
    st.stop()


# -----------------------------
# Process files
# -----------------------------
results = []
errors = []

with st.spinner("Analyzing PDF files..."):
    for file in uploaded_files:
        try:
            result = analyze_pdf(file.getvalue(), file.name, table_mode)
            results.append(result)
        except Exception as e:
            errors.append({"file_name": file.name, "error": str(e)})

if errors:
    st.error("Some files could not be processed.")
    st.dataframe(pd.DataFrame(errors), use_container_width=True)

if not results:
    st.stop()


# -----------------------------
# Overall overview
# -----------------------------
st.subheader("Batch overview")
overview_df = pd.DataFrame([
    {
        "file_name": r["file_name"],
        "pages": r["total_pages"],
        "text_extracted": r["has_text"],
        "tables_detected": r["total_tables"],
        "text_characters": len(r["full_text"]),
    }
    for r in results
])
st.dataframe(overview_df, use_container_width=True)

c1, c2, c3, c4 = st.columns(4)
c1.metric("Files processed", f"{len(results):,}")
c2.metric("Total pages", f"{int(overview_df['pages'].sum()):,}")
c3.metric("Files with text", f"{int(overview_df['text_extracted'].sum()):,}")
c4.metric("Tables detected", f"{int(overview_df['tables_detected'].sum()):,}")


# -----------------------------
# File explorer
# -----------------------------
st.subheader("File explorer")
file_names = [r["file_name"] for r in results]
selected_name = st.selectbox("Choose a PDF", file_names)
selected_result = next(r for r in results if r["file_name"] == selected_name)

st.markdown("### Page summary")
st.dataframe(selected_result["page_summary_df"], use_container_width=True)


tab1, tab2, tab3, tab4 = st.tabs(["Text preview", "Tables preview", "Single-file downloads", "Notes"])

with tab1:
    if selected_result["full_text"].strip():
        preview_chars = st.slider("Preview text length", 500, 20000, 4000, step=500)
        st.text_area(
            "Extracted text preview",
            value=selected_result["full_text"][:preview_chars],
            height=400,
        )
    else:
        st.warning("No selectable text was extracted. This PDF may be scanned or image-based.")

with tab2:
    if selected_result["tables"]:
        table_labels = [f"Page {t['page_number']} - Table {t['table_number']}" for t in selected_result["tables"]]
        selected_table_label = st.selectbox("Choose a detected table", table_labels)
        selected_table = selected_result["tables"][table_labels.index(selected_table_label)]
        st.dataframe(selected_table["dataframe"], use_container_width=True)
    else:
        st.info("No tables were detected in this PDF.")

with tab3:
    stem = Path(selected_result["file_name"]).stem

    if export_docx:
        docx_bytes = make_docx_bytes(selected_result["file_name"], selected_result["full_text"], selected_result["page_summary_df"])
        st.download_button(
            "Download Word (.docx)",
            data=docx_bytes,
            file_name=f"{stem}_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if export_excel:
        excel_bytes = make_excel_bytes(
            selected_result["file_name"],
            [t["dataframe"] for t in selected_result["tables"]],
            selected_result["page_summary_df"],
        )
        st.download_button(
            "Download Excel (.xlsx)",
            data=excel_bytes,
            file_name=f"{stem}_tables.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if export_txt:
        st.download_button(
            "Download text (.txt)",
            data=selected_result["full_text"] if selected_result["full_text"].strip() else "No text could be extracted.",
            file_name=f"{stem}_extracted_text.txt",
            mime="text/plain",
        )

    if export_csv:
        if selected_result["tables"]:
            all_tables_df = pd.concat([t["dataframe"] for t in selected_result["tables"]], ignore_index=True, sort=False)
            st.download_button(
                "Download all detected tables as one CSV",
                data=all_tables_df.to_csv(index=False),
                file_name=f"{stem}_all_tables.csv",
                mime="text/csv",
            )
        else:
            st.info("CSV export is available when at least one table is detected.")

with tab4:
    st.markdown(
        "- Word export is based on extracted text, so complex layout is simplified.\n"
        "- Excel and CSV exports are best for tables detected by the PDF parser.\n"
        "- Scanned PDFs may need OCR in a later version because image-only pages often contain no selectable text."
    )


# -----------------------------
# Batch downloads
# -----------------------------
st.subheader("Batch downloads")
zip_bytes = make_zip_bytes(
    results,
    include_docx=export_docx,
    include_excel=export_excel,
    include_csv=export_csv,
    include_txt=export_txt,
)

st.download_button(
    "Download all outputs as ZIP",
    data=zip_bytes,
    file_name="pdf_converter_outputs.zip",
    mime="application/zip",
)


# -----------------------------
# Requirements helper
# -----------------------------
with st.expander("Deployment notes", expanded=False):
    st.code(
        """streamlit\npandas\npdfplumber\npython-docx\nxlsxwriter""",
        language="text",
    )
